from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK


def add_toc(_paragraph):
    """插入目录字段的正确方法"""

    # 添加字段开始标记
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')

    # 添加字段指令
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'  # 包含1-3级标题，带超链接

    # 添加字段结束标记
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    _paragraph.add_run().add_break(WD_BREAK.LINE)
    _paragraph.add_run().element.extend([fld_char_begin, instr_text, fld_char_end])
    _paragraph.add_run().add_break(WD_BREAK.PAGE)


# 创建文档
doc = Document("a.docx")

for i, paragraph in enumerate(doc.paragraphs):
    if i == 0:
        # 添加文档内容（必须使用Heading样式）
        paragraph.add_run().add_break(WD_BREAK.PAGE)

    if i == 1:
        # 先插入目录
        new_para = paragraph.insert_paragraph_before("目录", style="Title")
        add_toc(new_para)

# 保存文档
doc.save('b.docx')
print("文档已生成，请用Word打开后按F9更新目录")
